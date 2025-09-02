from fastapi import FastAPI
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

app = FastAPI()

class CVData(BaseModel):
    nombre_perfil: str
    perfil: str
    herramientas: str
    experiencia: str
    preparacion: str
    formacion: str
    idiomas: str

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

@app.post("/chat")
async def chat(request: Request):
    body = await request.json()
    prompt = body.get("prompt", "")

    r = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json",
        },
        json={
            "model": "gpt-4o-mini",   # aquí pones el modelo que quieras
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
    )
    return r.json()
    
@app.post("/generar_cv")
def generar_cv(data: CVData):
    doc = Document("fromato1.docx")

    def reemplazar_texto(paragraphs, marcador, nuevo_texto):
        for p in paragraphs:
            if marcador in p.text:
                p.clear()
                run = p.add_run(nuevo_texto)
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    reemplazar_texto(doc.paragraphs, "{{ Nombre completo y perfil tecnico}}", data.nombre_perfil)
    reemplazar_texto(doc.paragraphs, "{{Perfil}}", data.perfil)
    reemplazar_texto(doc.paragraphs, "{{Herramientas tecnológicas}}", data.herramientas)
    reemplazar_texto(doc.paragraphs, "{{Experiencia laboral}}", data.experiencia)
    reemplazar_texto(doc.paragraphs, "{{Preparación académica}}", data.preparacion)
    reemplazar_texto(doc.paragraphs, "{{Formación adicional}}", data.formacion)
    reemplazar_texto(doc.paragraphs, "{{Idiomas}}", data.idiomas)

    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            if any(word in run.text for word in ["Empresa:", "Puesto:", "Fechas:", "Funciones:"]):
                run.bold = True

    output_file = "CV_HLS.docx"
    doc.save(output_file)

    return {"message": "CV generado con éxito", "archivo": output_file}
